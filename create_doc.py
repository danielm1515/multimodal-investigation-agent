"""
create_doc.py  —  generates the Hebrew Word document explaining the assignment.
Run once:  python create_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_rtl(paragraph):
    """Force RTL on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def para(doc, text, bold=False, size=11, color=None, italic=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'David'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = 'David'
    run.font.size = Pt(11)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E75B6')
        tc_pr.append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data rows
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.text = val
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if r_i % 2 == 1:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DEEAF1')
                tc_pr.append(shd)
    return table


# ── main document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.right_margin = Cm(2)
    section.left_margin  = Cm(2)
    section.top_margin   = Cm(2)
    section.bottom_margin = Cm(2)

# Default RTL for the whole document
doc.core_properties.language = 'he-IL'

# ─── TITLE PAGE ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('מטלה מעשית: סוכן AI מולטימודלי')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.font.name = 'David'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Multimodal Investigation Agent — Sense → Plan → Act → Observe')
run2.font.size = Pt(14)
run2.italic = True
run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('תיעוד מלא: מה בנינו, למה, ואיך — שלב אחר שלב')
run3.font.size = Pt(12)
run3.font.name = 'David'

doc.add_page_break()

# ─── 1. מבוא ─────────────────────────────────────────────────────────────────
heading(doc, '1. מבוא — מה המטרה?', 1, color=(0x1F, 0x4E, 0x79))

para(doc,
     'המטלה ביקשה מאיתנו לבנות סוכן AI מולטימודלי שעובד לפי מכונת מצבים (Finite State Machine). '
     'המשמעות היא שהסוכן לא מקבל קלט, שולח אותו ישירות למודל, ומחזיר תשובה — '
     'אלא עובר תהליך מורכב של חשיבה, תכנון, ביצוע, ובדיקה לפני כל פלט.')

para(doc,
     'הלולאה המרכזית שהגדרה המטלה:', bold=True)

code_block(doc, 'Sense → Plan → Act → Observe → Update State → Validate → Respond / Retry / Clarify')

para(doc,
     'הסוכן שבנינו קרוי Multimodal Investigation Agent — הוא מקבל "תיק חקירה" '
     'המורכב לפחות משתי מודליות (תמונה + מסמך, לדוגמה), מחלץ ראיות מכל אחת, '
     'מאחד אותן, ומחזיר תשובה מבוססת עם ציון אמינות (Confidence).')

doc.add_paragraph()

# ─── 2. ארכיטקטורה ───────────────────────────────────────────────────────────
heading(doc, '2. ארכיטקטורת המערכת', 1, color=(0x1F, 0x4E, 0x79))

para(doc,
     'הפרויקט מחולק לשבעה קבצי Python + קבצי עזר. כל קובץ אחראי על אחד מרכיבי הסוכן:',
     bold=False)

add_table(doc,
    ['קובץ', 'תפקיד'],
    [
        ['app.py',       'נקודת כניסה — טוען קלט, מריץ סוכן, מדפיס trace, שומר פלט'],
        ['agent.py',     'הסוכן עצמו — לולאת plan → act → observe'],
        ['state.py',     'יוצר אובייקט state חיצוני (זיכרון עבודה)'],
        ['planner.py',   'מחליט מהי הפעולה הבאה לפי control_state ו-guards'],
        ['tools.py',     'כלי הסוכן: analyze_image, analyze_document, transcribe_audio, generate_answer'],
        ['validator.py', 'בודק grounding ו-confidence לפני הפלט'],
        ['prompts.py',   'תבניות פרומפטים למודלים'],
    ]
)

doc.add_paragraph()

# ─── 3. מכונת המצבים ──────────────────────────────────────────────────────────
heading(doc, '3. מכונת המצבים (Finite State Machine)', 1, color=(0x1F, 0x4E, 0x79))

para(doc,
     'זוהי ליבת המטלה. הסוכן חייב לנוע בין מצבים מוגדרים בלבד — לא "לקפוץ" '
     'ישירות מקלט לתשובה. 12 המצבים המינימליים שנדרשו:')

add_table(doc,
    ['מצב', 'מה קורה כאן', 'המצב הבא'],
    [
        ['INGRESS',           'קלט משתמש מתקבל, מוודאים שיש קבצים',        'DETECT_MODALITIES'],
        ['DETECT_MODALITIES', 'מזהים אילו מודליות קיימות בקבצים',            'SELECT_TOOLS'],
        ['SELECT_TOOLS',      'בוחרים כלי ניתוח לפי המודליות שנמצאו',       'PLAN_NEXT_ACTION'],
        ['PLAN_NEXT_ACTION',  'מחליטים: האם לבצע כלי נוסף או לעבור לvalidation', 'ACT / VALIDATE'],
        ['ACT',               'מפעילים כלי אחד (תמונה / מסמך / אודיו)',      'OBSERVE'],
        ['OBSERVE',           'מקבלים את תוצאת הכלי ומעדכנים state',        'EXTRACT_EVIDENCE'],
        ['EXTRACT_EVIDENCE',  'שומרים את הראיה ברשימת evidence',             'PLAN_NEXT_ACTION'],
        ['VALIDATE',          'בודקים grounding: ≥2 מודליות + confidence',   'RESPOND / CLARIFY'],
        ['RESPOND',           'מייצרים תשובה סופית מבוססת ראיות',           'DONE'],
        ['CLARIFY',           'ראיות לא מספיקות — מבקשים הבהרה',            'DONE'],
        ['ERROR_RECOVERY',    'כלי נכשל — מנסה שוב (עד 2 פעמים) או fallback', 'PLAN_NEXT_ACTION'],
        ['DONE',              'הסוכן סיים',                                  '—'],
    ]
)

doc.add_paragraph()
para(doc,
     'הלולאה ACT→OBSERVE→EXTRACT_EVIDENCE→PLAN_NEXT_ACTION חוזרת על עצמה לכל קובץ. '
     'אחרי שכל הקבצים עובדו, הסוכן עובר ל-VALIDATE.',
     italic=True)

doc.add_paragraph()

# ─── 4. שלב אחר שלב ──────────────────────────────────────────────────────────
heading(doc, '4. שלב אחר שלב — מה בנינו ולמה', 1, color=(0x1F, 0x4E, 0x79))

# 4.1
heading(doc, '4.1  state.py — אובייקט State חיצוני', 2, color=(0x2E, 0x75, 0xB6))

para(doc,
     'למה? — המטלה אסרה להכניס הכול לפרומפט אחד. State חיצוני מאפשר לשמור את '
     'כל המידע בין שלבים ולשנות אותו בכל שלב בנפרד.')

para(doc, 'מה כלול ב-state:', bold=True)
for item in [
    'goal + success_criteria — מה הסוכן מנסה להשיג',
    'control_state — המצב הנוכחי ב-FSM',
    'available_modalities — אילו מודליות זוהו',
    'pending_tools — תור הכלים שעוד לא רצו',
    'evidence — רשימת הראיות שנאספו',
    'validation — תוצאת בדיקת הgrounding',
    'actions_taken — הטרייס המלא של כל שלב',
    'total_tokens — סה"כ tokens שנוצלו',
]:
    bullet(doc, item)

code_block(doc,
    'state = {\n'
    '  "goal": "Answer the user\'s question using grounded multimodal evidence",\n'
    '  "control_state": "INGRESS",\n'
    '  "available_modalities": [],\n'
    '  "pending_tools": [],\n'
    '  "evidence": [],\n'
    '  "validation": {"grounded": False, "confidence": 0.0, ...},\n'
    '  "total_tokens": {"prompt": 0, "completion": 0, "total": 0},\n'
    '  ...\n'
    '}'
)

doc.add_paragraph()

# 4.2
heading(doc, '4.2  tools.py — כלי הסוכן', 2, color=(0x2E, 0x75, 0xB6))

para(doc,
     'למה? — הסוכן חייב לממש לפחות שלושה כלים. כל כלי מחלץ ראיות ממודליות שונות.')

add_table(doc,
    ['כלי', 'מודליות', 'מודל אמיתי', 'Fallback'],
    [
        ['analyze_image',    'image',    'GPT-4o (vision)',    'mock תיאור'],
        ['analyze_document', 'document', 'GPT-4o-mini (text)', 'טקסט גולמי מהקובץ'],
        ['transcribe_audio', 'audio',    '— (mock בלבד)',      'תיאור mock'],
        ['generate_answer',  'כל הראיות', 'GPT-4o-mini',       'תשובה דטרמיניסטית'],
        ['validate_evidence','—',        'קוד לוגי',           '—'],
    ]
)

doc.add_paragraph()
para(doc,
     'חשוב: כל קריאת OpenAI מחזירה גם token usage. '
     'שמרנו את הנתונים הבאים לכל קריאה:', bold=False)

for item in [
    'prompt_tokens — כמה tokens נשלחו למודל',
    'completion_tokens — כמה tokens הגיעו בחזרה',
    'total_tokens — הסכום הכולל',
]:
    bullet(doc, item)

code_block(doc,
    '# דוגמה מ-analyze_image:\n'
    'tokens = {\n'
    '    "prompt":     resp.usage.prompt_tokens,\n'
    '    "completion": resp.usage.completion_tokens,\n'
    '    "total":      resp.usage.total_tokens,\n'
    '}\n'
    '# תוצאה בפועל (ריצה לדוגמה - המספרים משתנים מעט בין ריצות):\n'
    '# analyze_image    -> prompt: 857 | completion: 120 | total: 977\n'
    '# analyze_document -> prompt: 118 | completion: 103 | total: 221\n'
    '# generate_answer  -> prompt: 337 | completion: 127 | total: 464\n'
    '# -----------------------------------------------\n'
    '# TOTAL               prompt: 1312 | completion: 350 | total: 1662'
)

para(doc,
     'כל token נספר פעם אחת בלבד — בשלב ACT, שבו הכלי באמת קורא ל-OpenAI. '
     'הסכום הכולל נשמר ב-state["total_tokens"] ומוצג בסוף הריצה. '
     'המספרים המדויקים משתנים מעט בין ריצות כי אורך התשובה של המודל אינו דטרמיניסטי.',
     italic=True, size=10)

doc.add_paragraph()

# 4.3
heading(doc, '4.3  planner.py — המתכנן', 2, color=(0x2E, 0x75, 0xB6))

para(doc,
     'למה? — הסוכן לא קובע מה לעשות מתוך פרומפט. Planner מפורש קובע '
     'לפי control_state מה הפעולה הבאה, כולל guards (תנאי מעבר).')

para(doc, 'Guards מרכזיים:', bold=True)
for item in [
    'files_exist — INGRESS: אם אין קבצים → CLARIFY',
    'modalities_count ≥ 2 — SELECT_TOOLS: אם פחות משתי מודליות → CLARIFY',
    'pending_tools not empty — PLAN_NEXT_ACTION: האם יש עוד כלים לרוץ?',
    'retries < 2 — ERROR_RECOVERY: האם לנסות שוב או להשתמש ב-fallback?',
    'max_steps — שמירה מפני לולאה אינסופית',
]:
    bullet(doc, item)

code_block(doc,
    '# דוגמה: המתכנן ב-PLAN_NEXT_ACTION\n'
    'if state["pending_tools"]:\n'
    '    return {"action": "plan_next_action", "next_state": "ACT"}\n'
    'else:\n'
    '    return {"action": "plan_next_action", "next_state": "VALIDATE"}'
)

doc.add_paragraph()

# 4.4
heading(doc, '4.4  validator.py — בדיקת Grounding', 2, color=(0x2E, 0x75, 0xB6))

para(doc,
     'למה? — המטלה דרשה שהתשובה תהיה מבוססת על לפחות שתי מודליות. '
     'ה-Validator בודק זאת לפני שהסוכן מחזיר תשובה.')

para(doc, 'תנאי grounding:', bold=True)
for item in [
    'used_modalities ≥ 2 (דוגמה: image + document)',
    'average confidence ≥ minimum_confidence (ברירת מחדל: 0.5)',
    'evidence list אינה ריקה',
]:
    bullet(doc, item)

para(doc,
     'אם grounding = False → הסוכן עובר ל-CLARIFY ומבקש מידע נוסף במקום לנחש.')

doc.add_paragraph()

# 4.5
heading(doc, '4.5  agent.py — לולאת ה-FSM', 2, color=(0x2E, 0x75, 0xB6))

para(doc,
     'זהו הקובץ המרכזי. הוא מחבר בין כל הרכיבים ומריץ את הלולאה:')

code_block(doc,
    'while state["control_state"] != "DONE":\n'
    '    decision = self.plan()      # שואל את ה-planner\n'
    '    result   = self.act(decision)  # מבצע את הפעולה\n'
    '    self.observe(result)        # מעדכן state + trace'
)

para(doc, 'שלוש הפונקציות:', bold=True)
for item in [
    'plan()  — מפנה ל-planner, מקבל action + next_state',
    'act()   — מפעיל את הכלי המתאים לפי שם הפעולה',
    'observe() — מעדכן state, צובר tokens, מוסיף לטרייס, קובע next_state',
]:
    bullet(doc, item)

doc.add_paragraph()

# 4.6
heading(doc, '4.6  app.py — הפלט המפורט', 2, color=(0x2E, 0x75, 0xB6))

para(doc,
     'הוספנו פונקציה _print_trace() שמציגה לכל שלב:')

for item in [
    'מספר השלב, שם המצב, שם האירוע',
    'פירוט הכלי שרץ, המודליות, ה-source',
    'תצוגה מקדימה של התוכן שחולץ',
    'token usage לאותה קריאת API: prompt / completion / total',
    'מצטבר (cumulative) עד לאותו שלב',
    'סיכום כלל הtokens בסוף הריצה',
]:
    bullet(doc, item)

doc.add_paragraph()

# ─── 5. דוגמת הרצה מלאה ──────────────────────────────────────────────────────
heading(doc, '5. דוגמת הרצה — 14 שלבים', 1, color=(0x1F, 0x4E, 0x79))

para(doc,
     'הסוכן רץ על שני קבצים: examples/dashboard.png (תמונת גרף מכירות) '
     'ו-examples/context.txt (מסמך הסבר על בעיות אספקה).')

add_table(doc,
    ['שלב', 'מצב', 'אירוע', 'פרטים', 'Tokens'],
    [
        ['1',  'INGRESS',          'ingress_ok',          'קלט התקבל, יש קבצים',                       '—'],
        ['2',  'DETECT_MODALITIES','modalities_detected', 'זוהה: image + document',                     '—'],
        ['3',  'SELECT_TOOLS',     'tools_selected',      'analyze_image, analyze_document, generate_answer', '—'],
        ['4',  'PLAN_NEXT_ACTION', 'planned',             'תור ראשון: analyze_image על dashboard.png',  '—'],
        ['5',  'ACT',              'evidence',            'GPT-4o ניתח את התמונה',                      '977'],
        ['6',  'OBSERVE',          'observed',            'תוצאה נשמרה ב-last_tool_result',             '—'],
        ['7',  'EXTRACT_EVIDENCE', 'evidence_extracted',  'ראיה נוספה לרשימת evidence (modality=image)','—'],
        ['8',  'PLAN_NEXT_ACTION', 'planned',             'תור שני: analyze_document על context.txt',   '—'],
        ['9',  'ACT',              'evidence',            'GPT-4o-mini ניתח את המסמך',                  '221'],
        ['10', 'OBSERVE',          'observed',            'תוצאה נשמרה',                                '—'],
        ['11', 'EXTRACT_EVIDENCE', 'evidence_extracted',  'ראיה נוספה (modality=document)',             '—'],
        ['12', 'PLAN_NEXT_ACTION', 'planned',             'אין כלים נוספים → VALIDATE',                 '—'],
        ['13', 'VALIDATE',         'validation_result',   'grounded=True, confidence=0.81',             '—'],
        ['14', 'RESPOND→DONE',     'final_answer',        'GPT-4o-mini יצר תשובה סופית',               '464'],
    ]
)

para(doc, 'ה-tokens מדווחים בשלב ACT (שם הכלי קורא ל-OpenAI). סה"כ הריצה: 1,662 tokens '
          '(977 + 221 + 464). המספרים משתנים מעט בין ריצות.',
     italic=True, size=9)

doc.add_paragraph()
para(doc, 'תשובת הסוכן הסופית:', bold=True)
para(doc,
     '"הבעיה המרכזית היא ירידה חדה במכירות החל מאפריל, הקשורה לבעיות שרשרת אספקה. '
     'הצעד הבא: ניתוח מעמיק של ה-dashboard ביחס לנסיבות התפעוליות. '
     'Confidence: 0.81 | Used modalities: [document, image]"',
     italic=True, color=(0x1F, 0x4E, 0x79))

doc.add_paragraph()

# ─── 6. מדוע זה AI אמיתי ──────────────────────────────────────────────────────
heading(doc, '6. למה זה סוכן AI ולא סתם קוד?', 1, color=(0x1F, 0x4E, 0x79))

add_table(doc,
    ['תכונה', 'מה שיש כאן', 'מה שהיה מספיק (אבל לא בוצע)'],
    [
        ['Goal מפורש',      'אובייקט goal עם success_criteria',        'משימה בתוך פרומפט'],
        ['State חיצוני',    'dict שמתעדכן בין שלבים',                  'משתנים גלובליים'],
        ['Planner',         'קוד מפורש עם guards ו-transitions',       'if/else ישיר'],
        ['Tool Selection',  'ניתוב דינמי לפי מודליות שנמצאו',          'hardcode שם הכלי'],
        ['Observer',        'מעדכן state לפי סוג התוצאה',              'שמירת תוצאה בלבד'],
        ['Validator',       'בודק grounding לפני כל תשובה',            'תמיד מחזיר תשובה'],
        ['Error Recovery',  'retry + fallback + מניית ניסיונות',       'try/except פשוט'],
        ['Trace / Memory',  'היסטוריית פעולות מלאה + token tracking',  'הדפסה בלבד'],
    ]
)

doc.add_paragraph()

# ─── 7. סיכום ────────────────────────────────────────────────────────────────
heading(doc, '7. סיכום', 1, color=(0x1F, 0x4E, 0x79))

para(doc,
     'בנינו סוכן AI מולטימודלי שפועל לפי הכללים הבאים:')

for item in [
    'לא מחזיר תשובה לפני שבדק שיש ראיות מ-2 מודליות לפחות',
    'לא "מנחש" — כל משפט בתשובה מבוסס על evidence שחולץ בפועל מהקבצים',
    'עובר דרך 12 מצבים מוגדרים — לא קיצורי דרך',
    'מתאושש מכשלונות (retry + fallback) מבלי לקרוס',
    'שומר trace מלא + כמות tokens לכל שלב',
    'עובד הן עם OpenAI אמיתי והן ב-mock mode ללא API key',
]:
    bullet(doc, item)

doc.add_paragraph()
para(doc,
     'זהו בדיוק ההבדל בין "Upload → Model → Answer" לבין סוכן AI אמיתי: '
     'הסוכן חושב, מתכנן, מבצע, בודק — ורק אז עונה.',
     bold=True, color=(0x1F, 0x4E, 0x79))

# ─── save ─────────────────────────────────────────────────────────────────────
out = 'multimodal_agent_report.docx'
doc.save(out)
print(f"Saved: {out}")
